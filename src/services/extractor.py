"""
Document parser & structured lab-data extractor (Agent 1).

Extraction flow
================
  PDF   → pypdfium2 / pdfplumber (best-of-two) → LLM structured extraction
          ↳ scanned, or suspiciously few results → page images → Gemini Vision
  Image → Gemini Vision multimodal → structured JSON
  DOCX  → paragraphs + tables → LLM structured extraction
  Text  → direct LLM structured extraction
          ↳ LLM unavailable → regex fallback (clearly reported as degraded)

Accuracy notes
--------------
* **Structured output.** Gemini is called with a ``response_schema``, so the
  reply is schema-valid JSON by construction. The old free-form "return only
  JSON" prompt was the single largest source of dropped results.
* **Chunking, not truncation.** Long reports are split on page boundaries and
  extracted concurrently; previously everything past 15 000 characters was
  silently discarded.
* **Provenance.** Every run reports which engine and which model produced the
  results, so a degraded (regex) run is never mistaken for a good one.
"""

from __future__ import annotations

import asyncio
import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
from PIL import Image
from pydantic import BaseModel, Field

from config.settings import settings
from src.services.llm_factory import get_chat_model, get_vision_model
from src.services.pdf_text import extract_pdf_text, render_pdf_pages, repair_mojibake

logger = logging.getLogger(__name__)


# ───────────────────────── Tuning constants ─────────────────────────────────

CHUNK_CHAR_BUDGET = 12_000     # characters of report text per LLM call
MAX_CONCURRENT_CALLS = 4       # keep well inside free-tier rate limits
MAX_RETRIES = 3                # per LLM call, on transient errors
VISION_FALLBACK_MIN_ITEMS = 3  # below this, re-try a text PDF through vision


# ─────────────────────── Structured output schema ───────────────────────────
#
# Passed to Gemini as `response_schema`, which constrains decoding to valid
# JSON. Descriptions are part of the prompt the model sees — they carry real
# instruction weight, so they are written as directives.


class LabItem(BaseModel):
    """One row of a lab report."""

    test_name: str = Field(
        description="Test name exactly as printed, including any parenthesised "
        "abbreviation, e.g. 'Mean Corpuscular Volume (MCV)'."
    )
    observed_value: str = Field(
        description="The measured result verbatim: '14.5', 'Not seen', "
        "'Positive', '<0.5'. Never compute or infer a value."
    )
    unit: Optional[str] = Field(
        default=None,
        description="Unit of measurement, e.g. 'g/dL', 'mill/mm³', '10³/µL'. "
        "Often printed on the line BELOW the value. Null if the test is "
        "qualitative.",
    )
    reference_range: Optional[str] = Field(
        default=None,
        description="Biological reference interval as printed, e.g. '13 - 17', "
        "'< 200', 'Not seen'. Null if absent.",
    )
    panel: Optional[str] = Field(
        default=None,
        description="Section/panel heading this test sits under, e.g. "
        "'Complete Blood Count (CBC)', 'Differential Count', 'Lipid Profile'.",
    )
    method: Optional[str] = Field(
        default=None,
        description="Measurement method if printed, e.g. 'Flow Cytometry', "
        "'Calculated', 'DC Impedance Method'.",
    )
    flag: Optional[str] = Field(
        default=None,
        description="Abnormality marker as printed next to the value only "
        "('High', 'Low', 'H', 'L', 'Abnormal'). Null if the report prints none "
        "— do NOT derive it yourself.",
    )


class PatientInfo(BaseModel):
    """Demographics and report metadata."""

    name: Optional[str] = None
    age: Optional[str] = Field(default=None, description="e.g. '23 Year(s)' or '35'.")
    sex: Optional[str] = Field(default=None, description="'Male', 'Female', or null.")
    date: Optional[str] = Field(default=None, description="Report/reported-on date.")
    lab_name: Optional[str] = None
    referring_doctor: Optional[str] = Field(
        default=None, description="Referring physician ('Ref. by'), not the pathologist."
    )
    patient_id: Optional[str] = None
    report_ref_id: Optional[str] = None
    collected_on: Optional[str] = None
    reported_on: Optional[str] = None


class ExtractionResponse(BaseModel):
    """Top-level structured reply expected from the model."""

    patient_info: PatientInfo = Field(default_factory=PatientInfo)
    lab_results: List[LabItem] = Field(default_factory=list)
    notes: List[str] = Field(
        default_factory=list,
        description="Verbatim clinical remarks printed on the report, e.g. a "
        "\"Doctor's Note\" or interpretive comment. Exclude marketing copy, "
        "disclaimers and terms & conditions.",
    )


# ─────────────────────────── Prompt templates ───────────────────────────────

_SHARED_RULES = """\
─── RULES ───
• Extract EVERY test row, including normal ones. Missing a row is the worst
  possible error.
• Expand panels: if a panel such as CBC lists Hemoglobin, WBC, … emit one
  object per sub-test, and set `panel` to the panel heading.
• Copy `observed_value` verbatim. Never calculate, round, convert or infer.
• LAYOUT WARNING: in many reports the value and reference range sit on one
  line while the unit and the method sit on the lines directly above or
  below it. Example:
      Hemoglobin (Hb)
      Cyanide-free SLS method
      15.3 13 - 17
      g/dL
  This is ONE test: name 'Hemoglobin (Hb)', value '15.3', unit 'g/dL',
  reference_range '13 - 17', method 'Cyanide-free SLS method'.
• A test name may wrap across two lines — rejoin it before emitting.
• Qualitative results are tests too ('Not seen', 'Positive', 'Nil').
• IGNORE page headers/footers, addresses, phone numbers, marketing copy,
  terms & conditions, and the doctor's signature block. Do not emit these
  as tests.
• Set a field to null when the report does not print it. Never invent values.
"""

EXTRACTION_PROMPT = f"""\
You are a meticulous medical lab report parser. Extract all lab test results,
patient information and clinical notes from the report text below.

{_SHARED_RULES}
─── LAB REPORT TEXT ───
{{text}}
"""

VISION_PROMPT = f"""\
You are a meticulous medical lab report parser. Read this lab report page image
and extract all lab test results, patient information and clinical notes.

{_SHARED_RULES}
• Read values directly off the image. If a digit is genuinely unreadable, omit
  that row rather than guessing.
"""


# ────────────────────────── Result container ────────────────────────────────


@dataclass
class ExtractionResult:
    """Everything Agent 1 learned about a document."""

    items: List[Dict[str, Any]] = field(default_factory=list)
    patient_info: Optional[Dict[str, Any]] = None
    notes: List[str] = field(default_factory=list)
    method: str = "none"          # e.g. "gemini:text(pypdfium2)"
    degraded: bool = False        # True when we fell back to regex
    warnings: List[str] = field(default_factory=list)

    def as_tuple(self) -> Tuple[List[Dict[str, Any]], Optional[Dict[str, Any]]]:
        return self.items, self.patient_info


# ──────────────────────────── Public API ────────────────────────────────────


async def extract_from_file_detailed(file_path: str, file_type: str) -> ExtractionResult:
    """
    Extract lab results + patient info from a file, with provenance.

    Parameters
    ----------
    file_path : str
        Path to the uploaded file.
    file_type : str
        ``"pdf"`` | ``"image"`` | ``"docx"`` | ``"text"``
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if file_type == "pdf":
        return await _extract_from_pdf(path)
    if file_type == "image":
        return await _extract_from_image(path)
    if file_type == "docx":
        return await _extract_from_docx(path)
    if file_type == "text":
        return await _extract_from_text(path.read_text(encoding="utf-8", errors="replace"))

    raise ValueError(
        f"Unsupported file_type: '{file_type}'. Use 'pdf', 'image', 'docx', or 'text'."
    )


async def extract_from_file(
    file_path: str, file_type: str
) -> Tuple[List[Dict[str, Any]], Optional[Dict[str, Any]]]:
    """Backwards-compatible wrapper returning ``(items, patient_info)``."""
    result = await extract_from_file_detailed(file_path, file_type)
    return result.as_tuple()


def detect_file_type(filename: str) -> str:
    """Infer ``file_type`` from a filename extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}:
        return "image"
    if ext in {".docx", ".doc"}:
        return "docx"
    if ext in {".txt", ".csv", ".text", ".md"}:
        return "text"
    raise ValueError(f"Unsupported file extension: '{ext}'")


# ───────────────────────── PDF extraction ───────────────────────────────────


async def _extract_from_pdf(path: Path) -> ExtractionResult:
    """Parse a PDF; fall back to Gemini Vision for scans and thin results."""
    pdf = extract_pdf_text(path)

    if not pdf.is_usable:
        logger.info("PDF has no embedded text (%d chars) — using vision", pdf.char_count)
        result = await _extract_via_vision_pdf(path)
        result.warnings.append("PDF appears to be a scan; results came from image OCR.")
        return result

    result = await _extract_from_pages(pdf.pages, source=f"pdf:{pdf.engine}")

    # A text-bearing PDF that yields almost nothing usually means the text layer
    # is decorative (an image-backed PDF with a junk overlay) or the layout
    # defeated the parser. Vision is a genuinely different code path, so retry.
    if not result.degraded and len(result.items) < VISION_FALLBACK_MIN_ITEMS:
        logger.warning(
            "Only %d item(s) from the text layer — retrying via vision",
            len(result.items),
        )
        vision = await _extract_via_vision_pdf(path)
        if len(vision.items) > len(result.items):
            vision.warnings.append(
                f"Text layer yielded only {len(result.items)} result(s); "
                "used image-based extraction instead."
            )
            return vision

    return result


async def _extract_via_vision_pdf(path: Path) -> ExtractionResult:
    """Render each page and read it with Gemini Vision, in parallel."""
    try:
        images = render_pdf_pages(path)
    except Exception as exc:
        logger.error("Failed to render PDF pages: %s", exc)
        return ExtractionResult(method="vision:render-failed", warnings=[str(exc)])

    if not images:
        return ExtractionResult(method="vision:no-pages")

    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CALLS)

    async def one(index: int, img: Image.Image) -> Optional[ExtractionResponse]:
        async with semaphore:
            logger.info("Vision-extracting page %d/%d …", index + 1, len(images))
            return await _call_gemini_vision(img)

    responses = await asyncio.gather(
        *(one(i, img) for i, img in enumerate(images)), return_exceptions=True
    )
    return _merge_responses(responses, method=f"gemini-vision:{settings.GEMINI_MODEL}")


# ──────────────────────── Image extraction ──────────────────────────────────


async def _extract_from_image(path: Path) -> ExtractionResult:
    """Extract lab data from an image file via Gemini Vision."""
    img = Image.open(str(path))
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    response = await _call_gemini_vision(img)
    return _merge_responses([response], method=f"gemini-vision:{settings.GEMINI_MODEL}")


# ───────────────────────── DOCX extraction ──────────────────────────────────


async def _extract_from_docx(path: Path) -> ExtractionResult:
    """Extract lab data from a Word (.docx) file: paragraphs + tables."""
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        logger.error("Failed to open .docx file: %s", exc)
        return ExtractionResult(method="docx:open-failed", warnings=[str(exc)])

    paragraphs_text = "\n".join(p.text for p in document.paragraphs if p.text.strip())

    tables_text = ""
    table_direct_items: List[Dict[str, Any]] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            tables_text += " | ".join(cells) + "\n"
            # Standard 4-column lab table: [Test, Value, Unit, Range]
            if len(cells) >= 4 and cells[0] and cells[1] and not _is_header_cell(cells[0]):
                table_direct_items.append(
                    {
                        "test_name": cells[0],
                        "observed_value": cells[1],
                        "unit": _nullify(cells[2]),
                        "reference_range": _nullify(cells[3]),
                    }
                )

    combined = paragraphs_text.strip()
    if tables_text.strip():
        combined += "\n\n--- EXTRACTED TABLES ---\n" + tables_text.strip()

    if len(combined) <= 50:
        logger.warning("DOCX appears empty or too short (%d chars)", len(combined))
        return ExtractionResult(
            items=_clean_items(table_direct_items),
            patient_info=_regex_patient_info(combined),
            method="docx:tables-only",
            degraded=True,
            warnings=["DOCX contained almost no text; used raw table cells."],
        )

    result = await _extract_from_text(combined)
    # The LLM path can come back empty on an unusual layout; direct table cells
    # are a better answer than nothing.
    if not result.items and table_direct_items:
        logger.info("Falling back to %d direct DOCX table rows", len(table_direct_items))
        result.items = _clean_items(table_direct_items)
        result.patient_info = result.patient_info or _regex_patient_info(combined)
        result.degraded = True
        result.warnings.append("LLM returned no rows; used raw DOCX table cells.")
    return result


def _is_header_cell(value: str) -> bool:
    return value.strip().lower() in {
        "test name", "test", "investigation", "parameter", "name", "description",
    }


# ────────────────────── Text → LLM extraction ──────────────────────────────


async def _extract_from_text(raw_text: str) -> ExtractionResult:
    """Extract from a single blob of text (splitting it if it is long)."""
    return await _extract_from_pages([raw_text], source="text")


async def _extract_from_pages(pages: List[str], source: str) -> ExtractionResult:
    """
    Extract from page-segmented text.

    Pages are grouped into chunks under :data:`CHUNK_CHAR_BUDGET` and sent
    concurrently, then merged and de-duplicated. This replaces the old
    hard truncation, which silently dropped everything past 15 000 characters.
    """
    chunks = _chunk_pages(pages)
    full_text = "\n".join(pages)

    if not chunks:
        return ExtractionResult(method=f"{source}:empty")

    provider = settings.MODEL_PROVIDER.lower()
    semaphore = asyncio.Semaphore(MAX_CONCURRENT_CALLS)

    async def one(index: int, chunk: str):
        async with semaphore:
            if len(chunks) > 1:
                logger.info("Extracting chunk %d/%d (%d chars)", index + 1, len(chunks), len(chunk))
            prompt = EXTRACTION_PROMPT.format(text=chunk)
            if provider == "gemini":
                return await _call_gemini_text(prompt)
            return await _call_langchain_text(prompt)

    responses = await asyncio.gather(
        *(one(i, c) for i, c in enumerate(chunks)), return_exceptions=True
    )

    failures = [r for r in responses if isinstance(r, Exception) or r is None]
    result = _merge_responses(
        responses, method=f"{provider}:{_active_model_name()}:{source}"
    )

    # Every LLM call failed → the model is unreachable/misconfigured. Regex is a
    # poor substitute, so say so loudly instead of returning junk as "success".
    if not result.items and len(failures) == len(responses):
        logger.error("All %d LLM call(s) failed — using regex fallback", len(responses))
        detail = next(
            (str(r) for r in responses if isinstance(r, Exception)), "unknown error"
        )
        return ExtractionResult(
            items=_clean_items(_regex_fallback(full_text)),
            patient_info=_regex_patient_info(full_text),
            method=f"regex-fallback:{source}",
            degraded=True,
            warnings=[
                "AI extraction was unavailable, so results come from a basic "
                f"pattern matcher and are likely incomplete. Cause: {detail[:300]}"
            ],
        )

    if failures:
        result.warnings.append(
            f"{len(failures)} of {len(responses)} document chunks failed to "
            "extract; some results may be missing."
        )
    return result


def _chunk_pages(pages: List[str]) -> List[str]:
    """Group pages into prompt-sized chunks, splitting any oversized page."""
    chunks: List[str] = []
    current: List[str] = []
    size = 0

    for i, page in enumerate(pages):
        if not page.strip():
            continue
        labelled = f"--- PAGE {i + 1} ---\n{page}" if len(pages) > 1 else page

        # A single page bigger than the budget is split on line boundaries.
        if len(labelled) > CHUNK_CHAR_BUDGET:
            if current:
                chunks.append("\n\n".join(current))
                current, size = [], 0
            chunks.extend(_split_long_text(labelled))
            continue

        if size + len(labelled) > CHUNK_CHAR_BUDGET and current:
            chunks.append("\n\n".join(current))
            current, size = [], 0
        current.append(labelled)
        size += len(labelled)

    if current:
        chunks.append("\n\n".join(current))
    return chunks


def _split_long_text(text: str) -> List[str]:
    """Split an oversized block on line boundaries, with a small overlap."""
    lines = text.split("\n")
    chunks: List[str] = []
    current: List[str] = []
    size = 0
    for line in lines:
        if size + len(line) > CHUNK_CHAR_BUDGET and current:
            chunks.append("\n".join(current))
            # Overlap a few lines so a test split across the boundary survives.
            current, size = current[-4:], sum(len(x) for x in current[-4:])
        current.append(line)
        size += len(line) + 1
    if current:
        chunks.append("\n".join(current))
    return chunks


# ─────────────────────────── LLM callers ────────────────────────────────────


def _active_model_name() -> str:
    provider = settings.MODEL_PROVIDER.lower()
    return {
        "gemini": settings.GEMINI_MODEL,
        "ollama": settings.OLLAMA_MODEL,
        "groq": settings.GROQ_MODEL,
    }.get(provider, "unknown")


def _gemini_config():
    """Structured-output config shared by the text and vision paths."""
    from google.genai import types as genai_types

    return genai_types.GenerateContentConfig(
        temperature=0.0,  # deterministic transcription, not creative writing
        response_mime_type="application/json",
        response_schema=ExtractionResponse,
        # Nothing here needs tools; disabling AFC removes a spurious SDK warning.
        automatic_function_calling=genai_types.AutomaticFunctionCallingConfig(
            disable=True
        ),
    )


_TRANSIENT = ("429", "500", "502", "503", "504", "deadline", "timeout", "unavailable", "overloaded")


def _is_transient(exc: Exception) -> bool:
    message = str(exc).lower()
    return any(token in message for token in _TRANSIENT)


async def _with_retries(coro_factory, what: str):
    """Run ``coro_factory()`` with exponential backoff on transient errors."""
    last: Optional[Exception] = None
    for attempt in range(MAX_RETRIES):
        try:
            return await coro_factory()
        except Exception as exc:
            last = exc
            if not _is_transient(exc) or attempt == MAX_RETRIES - 1:
                raise
            delay = 2**attempt
            logger.warning(
                "%s failed (attempt %d/%d): %s — retrying in %ds",
                what, attempt + 1, MAX_RETRIES, str(exc)[:200], delay,
            )
            await asyncio.sleep(delay)
    raise last  # pragma: no cover - loop always returns or raises


async def _call_gemini_text(prompt: str) -> ExtractionResponse:
    """Structured text extraction through the google-genai async client."""
    client = get_vision_model()

    async def call():
        return await client.aio.models.generate_content(
            model=settings.GEMINI_MODEL,
            contents=prompt,
            config=_gemini_config(),
        )

    response = await _with_retries(call, "Gemini text extraction")
    return _parse_response(response)


async def _call_gemini_vision(img: Image.Image) -> ExtractionResponse:
    """Structured vision extraction for one page image."""
    client = get_vision_model()

    async def call():
        return await client.aio.models.generate_content(
            model=settings.GEMINI_MODEL,
            contents=[VISION_PROMPT, img],
            config=_gemini_config(),
        )

    response = await _with_retries(call, "Gemini vision extraction")
    return _parse_response(response)


async def _call_langchain_text(prompt: str) -> ExtractionResponse:
    """Ollama / Groq path via LangChain, with schema coaxed into the prompt."""
    schema_hint = (
        "\n\nReturn ONLY a JSON object shaped like:\n"
        '{"patient_info": {"name": null, "age": null, "sex": null, "date": null,'
        ' "lab_name": null, "referring_doctor": null, "patient_id": null,'
        ' "report_ref_id": null, "collected_on": null, "reported_on": null},'
        ' "lab_results": [{"test_name": "", "observed_value": "", "unit": null,'
        ' "reference_range": null, "panel": null, "method": null, "flag": null}],'
        ' "notes": []}\n'
        "No markdown fences, no commentary."
    )

    async def call():
        llm = get_chat_model()
        return await llm.ainvoke(prompt + schema_hint)

    response = await _with_retries(call, "LLM text extraction")
    content = getattr(response, "content", response)
    if isinstance(content, list):
        text = "".join(
            part if isinstance(part, str) else part.get("text", "") for part in content
        )
    else:
        text = str(content)
    return _parse_json_text(text)


# ───────────────────── Response parsing & merging ───────────────────────────


def _parse_response(response: Any) -> ExtractionResponse:
    """Prefer the SDK's parsed object; fall back to parsing the raw text."""
    parsed = getattr(response, "parsed", None)
    if isinstance(parsed, ExtractionResponse):
        return parsed
    if isinstance(parsed, dict):
        return ExtractionResponse.model_validate(parsed)
    return _parse_json_text(getattr(response, "text", "") or "")


def _parse_json_text(text: str) -> ExtractionResponse:
    """Robustly parse a JSON reply that may be fenced or wrapped in prose."""
    cleaned = re.sub(r"^```(?:json)?\s*", "", text.strip())
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()

    data: Any = None
    try:
        data = json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError:
                pass

    if not isinstance(data, dict):
        logger.error("Could not parse JSON from LLM response (%d chars)", len(text))
        return ExtractionResponse()

    try:
        return ExtractionResponse.model_validate(data)
    except Exception:
        # Salvage whatever validates rather than losing the whole chunk.
        items = [
            LabItem.model_validate(i)
            for i in data.get("lab_results", [])
            if isinstance(i, dict) and i.get("test_name")
        ]
        info = data.get("patient_info")
        return ExtractionResponse(
            patient_info=PatientInfo.model_validate(info) if isinstance(info, dict) else PatientInfo(),
            lab_results=items,
            notes=[str(n) for n in data.get("notes", []) if n],
        )


def _merge_responses(responses: List[Any], method: str) -> ExtractionResult:
    """Combine per-chunk / per-page responses into one de-duplicated result."""
    items: List[Dict[str, Any]] = []
    notes: List[str] = []
    merged_info: Dict[str, Any] = {}

    for response in responses:
        if isinstance(response, Exception) or response is None:
            if isinstance(response, Exception):
                logger.error("Chunk extraction failed: %s", str(response)[:300])
            continue
        items.extend(item.model_dump() for item in response.lab_results)
        notes.extend(response.notes)
        # First non-null wins: page 1 carries the header block, later pages
        # repeat it, and a later page must not blank out an earlier value.
        for key, value in response.patient_info.model_dump().items():
            if merged_info.get(key) in (None, "") and _nullify(value):
                merged_info[key] = _nullify(value)

    clean = _clean_items(items)
    logger.info("Extracted %d lab item(s) via %s", len(clean), method)

    return ExtractionResult(
        items=clean,
        patient_info=merged_info or None,
        notes=_dedupe_strings(notes),
        method=method,
    )


# ────────────────────── Post-processing / validation ────────────────────────

# Lines the model occasionally mistakes for tests despite the prompt.
_NON_TEST_PATTERNS = re.compile(
    r"^(test|result|investigation|parameter|biological reference|reference "
    r"interval|approved by|reviewed by|page \d|patient (id|name)|report ref|"
    r"collected|received|reported|partner|ref\.? by|address|phone|email|"
    r"clinical significance|conditions of|disclaimer)\b",
    re.I,
)

_UNIT_CANON = {
    "mill/mm3": "mill/mm³",
    "cells/mm3": "cells/mm³",
    "/mm3": "/mm³",
    "10^3/ul": "10³/µL",
    "10*3/ul": "10³/µL",
    "ul": "µL",
    "ug/dl": "µg/dL",
    "uiu/ml": "µIU/mL",
}


def _clean_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Normalise, drop non-tests, and de-duplicate extracted rows."""
    clean: List[Dict[str, Any]] = []
    seen: set[tuple] = set()

    for item in items:
        if not isinstance(item, dict):
            continue
        name = repair_mojibake(str(item.get("test_name") or "")).strip().rstrip(":").strip()
        value = repair_mojibake(str(item.get("observed_value") or "")).strip()

        if not name or not value:
            continue
        if len(name) < 2 or _NON_TEST_PATTERNS.match(name):
            logger.debug("Dropping non-test row: %r", name)
            continue

        # Reference ranges legitimately span two lines (e.g. Mentzer Index);
        # collapse them so downstream consumers see a single-line string.
        ref = _nullify(item.get("reference_range"))
        if ref:
            ref = re.sub(r"\s*\n\s*", "; ", repair_mojibake(ref)).strip()

        record = {
            "test_name": name,
            "observed_value": value,
            "unit": _canon_unit(item.get("unit")),
            "reference_range": ref,
            "panel": _nullify(item.get("panel")),
            "method": _nullify(item.get("method")),
            "flag": _nullify(item.get("flag")),
        }

        # Chunk overlap and repeated page headers produce genuine duplicates.
        key = (name.lower(), value.lower())
        if key in seen:
            continue
        seen.add(key)
        clean.append(record)

    return clean


def _canon_unit(unit: Any) -> Optional[str]:
    value = _nullify(unit)
    if value is None:
        return None
    value = repair_mojibake(value).strip()
    return _UNIT_CANON.get(value.lower(), value)


def _dedupe_strings(values: List[str]) -> List[str]:
    seen: set[str] = set()
    out: List[str] = []
    for value in values:
        text = str(value).strip()
        if text and text.lower() not in seen:
            seen.add(text.lower())
            out.append(text)
    return out


# ────────────────────── Regex fallback extractor ────────────────────────────
#
# Only reached when the LLM is unreachable. It handles the two most common
# machine-readable layouts and is always reported as a degraded result.

_LAB_LINE = re.compile(
    r"^"
    r"(?P<name>.{3,45}?)"
    r"\s{2,}"
    r"(?P<value>[\d.,]+|[A-Za-z]+(?:\s[A-Za-z]+)?)"
    r"\s+"
    r"(?P<unit>[a-zA-Z/%µ·^0-9³²]+(?:/[a-zA-Z]+)?)"
    r"\s+"
    r"(?P<range>[\d.,]+\s*[-–—]\s*[\d.,]+|[<>≤≥]\s*[\d.,]+)"
    r"\s*$",
    re.MULTILINE,
)

_PIPE_LINE = re.compile(
    r"^\s*(?P<name>[^|\n\r]+?)\s*\|\s*(?P<value>[^|\n\r]+?)\s*"
    r"(?:\|\s*(?P<unit>[^|\n\r]*?)\s*)?(?:\|\s*(?P<range>[^|\n\r]*?)\s*)?$",
    re.MULTILINE,
)

# Stacked layout, common in PDF-native reports: the test name (sometimes
# wrapped over two lines) and its method sit above a "value  low - high" line,
# with the unit alone underneath:
#
#     Mean Corpuscular Hemoglobin
#     Concentration (MCHC)
#     Calculated
#     34.5 31.5 - 34.5
#     g/dL
#
# This is what defeated the old fallback on the Orange Health report. A single
# regex cannot tell a wrapped name from a method line, so the anchor line is
# matched with a regex and the surrounding lines are resolved in Python.
_VALUE_RANGE_LINE = re.compile(
    r"^(?P<value>[\d.]+)\s+(?P<range>[\d.]+\s*[-–—]\s*[\d.]+|[<>≤≥]=?\s*[\d.]+)\s*$"
)

_METHOD_LINE = re.compile(
    r"^(?:calculated|index value|ratio|"
    r".*\b(?:method|cytometry|impedance|stain|microscopy|photometry|elisa|"
    r"clia|hplc|nephelometry|spectrophotometry|electrode|assay)\b.*)$",
    re.I,
)

# A line that is only a unit — it belongs to the test above, never to the one
# below, so it must not be mistaken for the start of a name.
_UNIT_ONLY_LINE = re.compile(
    r"^(?:%|ratio|index value|[a-zA-Zµ³²]{1,6}(?:/[a-zA-Zµ³²]{1,6})?|"
    r"\d+[³²]/[a-zA-Zµ]{1,3}|/[a-zA-Zµ³²]{1,6}|[a-zA-Z]+/[a-zA-Z]+[³²])$",
    re.I,
)


def _extract_stacked_layout(raw_text: str) -> List[Dict[str, Any]]:
    """Recover tests from the stacked name / method / value / unit layout."""
    lines = [ln.strip() for ln in raw_text.split("\n")]
    found: List[Dict[str, Any]] = []
    consumed: set[int] = set()

    for i, line in enumerate(lines):
        m = _VALUE_RANGE_LINE.match(line)
        if not m:
            continue

        # Walk backwards over at most three lines to collect the method and the
        # (possibly wrapped) test name.
        method: Optional[str] = None
        name_parts: List[str] = []
        for j in range(i - 1, max(i - 4, -1), -1):
            candidate = lines[j]
            if j in consumed or not candidate:
                break
            if _UNIT_ONLY_LINE.match(candidate) or _NON_TEST_PATTERNS.match(candidate):
                break
            if not name_parts and _METHOD_LINE.match(candidate):
                method = candidate
                consumed.add(j)
                continue
            name_parts.insert(0, candidate)
            consumed.add(j)
            # A line ending in ')' or a long line is usually the whole name;
            # only continue collecting when the text clearly wrapped.
            if len(name_parts) >= 2 or candidate[0].isupper():
                break

        if not name_parts:
            continue

        unit = None
        if i + 1 < len(lines) and lines[i + 1] and not _VALUE_RANGE_LINE.match(lines[i + 1]):
            nxt = lines[i + 1]
            if len(nxt) <= 20 and not _NON_TEST_PATTERNS.match(nxt):
                unit = nxt
                consumed.add(i + 1)

        consumed.add(i)
        found.append(
            {
                "test_name": " ".join(name_parts),
                "observed_value": m.group("value"),
                "unit": unit,
                "reference_range": m.group("range"),
                "method": method,
            }
        )
    return found


def _regex_fallback(raw_text: str) -> List[Dict[str, Any]]:
    """Best-effort pattern extractor used only when the LLM is unavailable."""
    results: List[Dict[str, Any]] = []
    seen: set[str] = set()

    def add(name: str, value: str, unit: Any, ref: Any, method: Any = None):
        name = name.strip().rstrip(":").strip()
        if not name or name.lower() in seen or _NON_TEST_PATTERNS.match(name):
            return
        seen.add(name.lower())
        results.append(
            {
                "test_name": name,
                "observed_value": value.strip(),
                "unit": _nullify(unit),
                "reference_range": _nullify(ref),
                "method": _nullify(method),
                "panel": None,
                "flag": None,
            }
        )

    for row in _extract_stacked_layout(raw_text):
        add(
            row["test_name"],
            row["observed_value"],
            row["unit"],
            row["reference_range"],
            row["method"],
        )

    for m in _LAB_LINE.finditer(raw_text):
        add(m.group("name"), m.group("value"), m.group("unit"), m.group("range"))

    if not results:
        for m in _PIPE_LINE.finditer(raw_text):
            name, value = m.group("name").strip(), m.group("value").strip()
            if not name or not value or _is_header_cell(name):
                continue
            if name.startswith("---"):
                continue
            add(name, value, m.group("unit"), m.group("range"))

    logger.info("Regex fallback extracted %d item(s)", len(results))
    return results


_PAT_NAME = re.compile(r"(?:patient\s*name|name)\s*[:\-]\s*(.+?)(?:\n|$)", re.I)
_PAT_AGE = re.compile(r"(?:age)\s*[:\-]?\s*(\d{1,3})\s*(?:years?\(?s?\)?|yrs?)", re.I)
_PAT_AGE_INLINE = re.compile(r"\b(\d{1,3})\s*Year\(s\)", re.I)
_PAT_SEX = re.compile(r"(?:sex|gender)\s*[:\-]\s*(male|female|m|f)\b", re.I)
_PAT_SEX_INLINE = re.compile(r"Year\(s\)\s*/?\s*(Male|Female)\b", re.I)
_PAT_PID = re.compile(r"Patient\s*ID\s*[:\-]\s*(\S+)", re.I)
_PAT_REF = re.compile(r"Report\s*Ref\.?\s*ID\s*[:\-]\s*(\S+)", re.I)
_PAT_COLLECTED = re.compile(r"Collected\s*[:\-]\s*(.+?)(?:\n|$)", re.I)
_PAT_REPORTED = re.compile(r"Reported\s*(?:on)?\s*[:\-]\s*(.+?)(?:\n|$)", re.I)


def _regex_patient_info(raw_text: str) -> Optional[Dict[str, Any]]:
    """Extract patient demographics with simple patterns (degraded path)."""
    info: Dict[str, Any] = {}

    for key, pattern in (
        ("name", _PAT_NAME),
        ("patient_id", _PAT_PID),
        ("report_ref_id", _PAT_REF),
        ("collected_on", _PAT_COLLECTED),
        ("reported_on", _PAT_REPORTED),
    ):
        m = pattern.search(raw_text)
        if m:
            info[key] = m.group(1).strip()

    m = _PAT_AGE.search(raw_text) or _PAT_AGE_INLINE.search(raw_text)
    if m:
        info["age"] = m.group(1).strip()

    m = _PAT_SEX.search(raw_text) or _PAT_SEX_INLINE.search(raw_text)
    if m:
        raw = m.group(1).strip().upper()
        info["sex"] = "Male" if raw in ("M", "MALE") else "Female"

    return {k: v for k, v in info.items() if v} or None


# ──────────────────────────── Helpers ───────────────────────────────────────


def _nullify(value: Any) -> Optional[str]:
    """Convert empty / sentinel strings to ``None``."""
    if value is None:
        return None
    s = str(value).strip()
    if s.lower() in ("", "null", "none", "n/a", "na", "-", "--", "nil.", "not provided"):
        return None
    return s
